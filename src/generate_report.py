import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('FRUIT RUSH: AI Pathfinding & ML Prediction\nProject Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Team Info
    doc.add_paragraph("")
    team_p = doc.add_paragraph()
    team_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    team_p.add_run('Submitted By:\n').bold = True
    team_p.add_run('Ahsan Faizan (23k-0615)\n')
    team_p.add_run('Ahzam Hassan (23k-0695)\n')
    doc.add_paragraph("")

    # Section 1: Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "This project implements an intelligent autonomous navigation system using pathfinding "
        "algorithms paired with a Machine Learning classifier to predict the complexity of a given map layout. "
        "The project is packaged into a modern, interactive Pygame application where the algorithms can "
        "be visualized in different environments, including constraint-based pathfinding and real-time AI chasing."
    )

    # Section 2: Informed Search Algorithm (A*)
    doc.add_heading('2. Informed Search Algorithms', level=1)
    doc.add_paragraph(
        "The core logic of the navigation system depends on the A* (A-Star) search algorithm using the Manhattan "
        "Distance heuristic. It is highly optimized for grid-based maneuvering."
    )
    doc.add_heading('2.1 Standard A*', level=2)
    doc.add_paragraph(
        "Used for direct Start-to-Goal pathing. It finds the shortest available route while avoiding generated "
        "walls and obstacles."
    )
    doc.add_heading('2.2 Constraint-Based A*', level=2)
    doc.add_paragraph(
        "An extension of the algorithm where the state representation includes a structure of currently collected "
        "items. The agent must collect all spawned fruits on the map before advancing to the final goal. The "
        "implementation successfully avoids infinite loops and maps the shortest sequence correctly."
    )

    # Section 3: Machine Learning Model
    doc.add_heading('3. Machine Learning Approach', level=1)
    doc.add_paragraph(
        "To predict map complexity, a Decision Tree Classifier was trained. A random map generator was utilized "
        "to run standard A* across 1,000 synthetic random mazes. Feature extraction targeted three primary parameters:"
    )
    ul1 = doc.add_paragraph(style='List Bullet')
    ul1.add_run("Obstacle Density: ").bold = True
    ul1.add_run("The percentage of grid space blocked by walls.")
    
    ul2 = doc.add_paragraph(style='List Bullet')
    ul2.add_run("Euclidean Distance: ").bold = True
    ul2.add_run("The straight-line heuristic from start to goal.")
    
    ul3 = doc.add_paragraph(style='List Bullet')
    ul3.add_run("Turning Points: ").bold = True
    ul3.add_run("The amount of corners navigated during an optimal A* path.")

    doc.add_paragraph(
        "Paths requiring steps above the overall median were labeled 'Complex', and below were labeled 'Simple'. "
        "The Decision Tree achieved a 90.00% accuracy score on the test data. The resulting model is serialized "
        "and runs real-time inference within the game's UI."
    )

    # Section 4: Game Modes and Interactive Visuals
    doc.add_heading('4. Interactive Environment & Game Modes', level=1)
    doc.add_paragraph(
        "The interface was meticulously styled with a modern dark theme and neon accents. It features smooth "
        "animations and three main functional sections:"
    )
    doc.add_heading('4.1 Mode 1: Constraint A* Visualization', level=2)
    doc.add_paragraph(
        "The system visually simulates the A* algorithm exploring the grid in real-time using neon teal search "
        "blocks. Once the optimal path tracing through all fruits is found, the agent traverses the route smoothly."
    )
    doc.add_heading('4.2 Mode 2: AI Survival Chase', level=2)
    doc.add_paragraph(
        "Players manually navigate the grid using arrow keys to collect fruits while the AI progressively hunts "
        "them using dynamic A* re-pathing. The difficulty scales up continuously across levels, forcing the player "
        "to outsmart the AI."
    )
    doc.add_heading('4.3 Custom Map Editor', level=2)
    doc.add_paragraph(
        "Users can draw their own walls, spawn items, and dictate the start/goal nodes. The Machine Learning "
        "classifier instantly updates its complexity prediction upon any structural change."
    )

    # Conclusion
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        "The project successfully validates the efficiency of A* in heavily constrained environments, while "
        "simultaneously mapping out the geometric difficulty metrics that effectively classify a maze using "
        "Machine Learning. The application executes stably without graphical glitches and encapsulates "
        "advanced AI routines interactively."
    )

    output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "Fruit_Rush_Report.docx")
    doc.save(output_path)
    print(f"Report successfully saved to {output_path}")

if __name__ == '__main__':
    create_report()